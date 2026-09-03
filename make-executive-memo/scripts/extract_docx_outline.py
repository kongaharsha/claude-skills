#!/usr/bin/env python
"""Extract a lightweight outline from a DOCX memo.

Usage:
  python extract_docx_outline.py path/to/memo.docx --max-paragraphs 80 --tables 2
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def compact(text: str) -> str:
    return " ".join(text.split())


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", help="Path to the DOCX file")
    parser.add_argument("--max-paragraphs", type=int, default=80)
    parser.add_argument("--tables", type=int, default=2)
    parser.add_argument("--cell-width", type=int, default=140)
    args = parser.parse_args()

    path = Path(args.docx_path)
    doc = Document(str(path))

    print(f"FILE: {path}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}")
    print(f"TABLES: {len(doc.tables)}")
    print()

    count = 0
    for para in doc.paragraphs:
        text = compact(para.text)
        if not text:
            continue
        count += 1
        print(f"{count:02d}: {text}")
        if count >= args.max_paragraphs:
            break

    for table_index, table in enumerate(doc.tables[: args.tables], start=1):
        print()
        print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
        for row in table.rows[:8]:
            cells = [compact(cell.text)[: args.cell_width] for cell in row.cells]
            print(" | ".join(cells))


if __name__ == "__main__":
    main()
